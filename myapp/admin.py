from django.contrib import admin
from django.http import HttpResponse
from .models import FormData

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def generate_docx(modeladmin, request, queryset):
    for form_data in queryset:
        document = Document("C:\Projects\myproject\myapp\Zgloszenie_szablon.docx")
        table1 = document.tables[0]
        table2 = document.tables[1]
        # Tabela 1
        table1.cell(0, 3).paragraphs[0].text = form_data.nr_zgloszenia or ""
        table1.cell(0, 3).add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        table1.cell(0, 3).paragraphs[1].text = form_data.nr_EZD or ""
        if str(form_data.data_zgloszenia) == "None":
            table1.cell(1, 3).paragraphs[-1].text = ""
        else:
            table1.cell(1, 3).paragraphs[-1].text = str(form_data.data_zgloszenia)
        # Paragrafy
        document.paragraphs[2].add_run(form_data.nazwa_zakladu or "")
        document.paragraphs[3].add_run(form_data.laboratorium or "")
        document.paragraphs[4].add_run(form_data.zglaszajacy or "")
        document.paragraphs[5].add_run(form_data.telefon or "")
        document.paragraphs[6].add_run(form_data.nr_pomieszczenia or "")

        # Tabela 2
        table2.cell(0, 0).paragraphs[-1].add_run(form_data.nazwa_urzadzenia) or ""
        table2.cell(1, 0).paragraphs[-1].add_run(form_data.dostep_do_sieci or "")
        table2.cell(2, 0).paragraphs[-1].add_run(form_data.nr_ewidencyjny or "")
        table2.cell(2, 2).paragraphs[-1].add_run(form_data.nr_gniazdka_lan or "")
        table2.cell(3, 0).paragraphs[-1].add_run(form_data.opis_zgloszenia or "")
        table2.cell(4, 0).paragraphs[-1].add_run(form_data.oczekiwania or "")
        table2.cell(5, 3).text = form_data.istotnosc_pouf or ""
        table2.cell(6, 3).text = form_data.istotnosc_integr or ""
        table2.cell(7, 3).text = form_data.istotnosc_dost or ""
        table2.cell(9, 0).paragraphs[0].text = form_data.zglaszajacy_podpis or ""
        table2.cell(10, 0).paragraphs[1].add_run(form_data.kierownik_lim_opinia or "")
        table2.cell(11, 0).paragraphs[0].text = form_data.kierownik_lim_podpis or ""
        if str(form_data.kierownik_lim_data) == "None":
            table2.cell(11, 0).paragraphs[2].text = ""
        else:
            table2.cell(11, 0).paragraphs[2].text = str(form_data.kierownik_lim_data)

        table2.cell(12, 0).paragraphs[1].add_run(form_data.kierownik_km_opinia)
        table2.cell(13, 0).paragraphs[0].add_run(form_data.kierownik_km_podpis or "")
        if str(form_data.kierownik_km_data) == "None":
            table2.cell(13, 0).paragraphs[-1].text = ""
        else:
            table2.cell(13, 0).paragraphs[-1].text = str(form_data.kierownik_km_data)
        table2.cell(14, 0).paragraphs[1].add_run(form_data.realizacja_opis or "")
        table2.cell(15, 0).paragraphs[-1].text = form_data.realizacja_podpis or ""
        if str(form_data.realizacja_data) == "None":
            table2.cell(15, 0).paragraphs[-1].text = ""
        else:
            table2.cell(15, 0).paragraphs[-1].text = str(form_data.realizacja_data)

        table2.cell(16, 0).paragraphs[2].add_run(form_data.potwierdzenie_podpis or "")
        if str(form_data.potwierdzenie_data) == "None":
            table2.cell(16, 0).paragraphs[-1].text = ""
        else:
            table2.cell(16, 0).paragraphs[-1].text = str(form_data.potwierdzenie_data)

        # Clean the filename to avoid illegal characters
        cleaned_name = re.sub(r'[\\/*?:"<>|]', "", form_data.nr_zgloszenia)
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={cleaned_name}.docx'
        document.save(response)
        return response


generate_docx.short_description = "Generate DOCX for selected entries"


@admin.register(FormData)
class FormDataAdmin(admin.ModelAdmin):
    list_display = ('nr_zgloszenia', 'data_zgloszenia', 'nazwa_zakladu', 'laboratorium')
    search_fields = ('nr_zgloszenia', 'data_zgloszenia', 'nazwa_zakladu', 'laboratorium')
    list_filter = ('data_zgloszenia', 'nazwa_zakladu', 'laboratorium')
    ordering = ()
    actions = [generate_docx]
