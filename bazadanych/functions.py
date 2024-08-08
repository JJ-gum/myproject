from django.contrib import admin
from django.template.response import TemplateResponse
from django.urls import reverse
from django.shortcuts import redirect
from .models import Urzadzenie, Zgloszenie
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import pandas as pd
import unicodecsv as csv
from django.http import HttpResponse
from datetime import datetime
from django.db import connection
from io import BytesIO


def export_urzadzenie_to_csv(modeladmin, request, queryset):
    response = HttpResponse(content_type='text/csv')
    now = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    filename = f'urzadzenia_{now}.csv'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    writer = csv.writer(response, quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow([
        'PIM ID', 'Laboratorium', 'Nr Pomieszczenia', 'Opis', 'Nr Ewidencyjny', 'Typ Urządzenia',
        'System Operacyjny', 'CPU', 'RAM', 'Pamięć Dysku', 'Dodatkowe Komponenty', 'Oprogramowanie Specjalne',
        'Konta', 'Data Rejestracji', 'Nr Gniazda', 'Typ Gniazda', 'Opiekun 1', 'Opiekun 2',
        'Typ Połączenia Sieciowego', 'Notatki'
    ])

    for urzadzenie in queryset:
        system_operacyjny = ', '.join([sys.typ_system_operacyjny.replace('\n', ' ').replace('\r', ' ') for sys in urzadzenie.system_operacyjny.all()])

        def sanitize_field(field):
            if field is None:
                return ''
            return str(field).replace('\n', ' ').replace('\r', ' ')
        writer.writerow([
            sanitize_field(urzadzenie.pim_id),
            sanitize_field(urzadzenie.laboratorium),
            sanitize_field(urzadzenie.nr_pomieszczenia),
            sanitize_field(urzadzenie.opis),
            sanitize_field(urzadzenie.numer_ewidencyjny),
            sanitize_field(urzadzenie.typ_urzadzenia),
            sanitize_field(system_operacyjny),
            sanitize_field(urzadzenie.cpu),
            sanitize_field(urzadzenie.ram),
            sanitize_field(urzadzenie.pamiec_dysku),
            sanitize_field(urzadzenie.dodatkowe_komponenty),
            sanitize_field(urzadzenie.oprogramowanie_specjalne),
            sanitize_field(urzadzenie.konta),
            sanitize_field(urzadzenie.data_rejestracji),
            sanitize_field(urzadzenie.nr_gniazda),
            sanitize_field(urzadzenie.typ_gniazd),
            sanitize_field(urzadzenie.opiekun_1),
            sanitize_field(urzadzenie.opiekun_2),
            sanitize_field(urzadzenie.typ_polaczenia_sieciowego),
            sanitize_field(urzadzenie.notatki)
        ])
    return response


def export_zgloszenie_to_csv(modeladmin, request, queryset):
    response = HttpResponse(content_type='text/csv')
    now = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    filename = f'zgloszenia_{now}.csv'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    writer = csv.writer(response, quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow([
        'ID', 'EZD ID Koszulki', 'Nr Zgłoszenia', 'Nr EZD', 'Data Zgłoszenia', 'Nazwa Zakładu', 'Laboratorium',
        'Zgłaszający', 'Telefon', 'Nr Pomieszczenia', 'Nazwa Urządzenia', 'Dostęp do Sieci', 'Nr Ewidencyjny',
        'Nr Gniazdka LAN', 'Opis Zgłoszenia', 'Oczekiwania', 'Istotność Poufności', 'Istotność Integralności',
        'Istotność Dostępności', 'Zgłaszający Podpis', 'Kierownik LIM Opinia', 'Kierownik LIM Podpis',
        'Kierownik LIM Data', 'Kierownik KM Opinia', 'Kierownik KM Podpis', 'Kierownik KM Data',
        'Realizacja Opis', 'Realizacja Podpis', 'Realizacja Data', 'Potwierdzenie Podpis', 'Potwierdzenie Data',
        'PIM ID Urządzenia', 'Notatki'
    ])

    def sanitize_field(field):
        if field is None:
            return ''
        return str(field).replace('\n', ' ').replace('\r', ' ')

    for zgloszenie in queryset:
        urzadzenie_id = sanitize_field(zgloszenie.urzadzenie_id.pim_id) if zgloszenie.urzadzenie_id else ''
        writer.writerow([
            zgloszenie.id,
            sanitize_field(zgloszenie.nr_EZD_ID_koszulki),
            sanitize_field(zgloszenie.nr_zgloszenia),
            sanitize_field(zgloszenie.nr_EZD),
            sanitize_field(zgloszenie.data_zgloszenia),
            sanitize_field(zgloszenie.nazwa_zakladu),
            sanitize_field(zgloszenie.laboratorium),
            sanitize_field(zgloszenie.zglaszajacy),
            sanitize_field(zgloszenie.telefon),
            sanitize_field(zgloszenie.nr_pomieszczenia),
            sanitize_field(zgloszenie.nazwa_urzadzenia),
            sanitize_field(zgloszenie.dostep_do_sieci),
            sanitize_field(zgloszenie.nr_ewidencyjny),
            sanitize_field(zgloszenie.nr_gniazdka_lan),
            sanitize_field(zgloszenie.opis_zgloszenia),
            sanitize_field(zgloszenie.oczekiwania),
            sanitize_field(zgloszenie.istotnosc_pouf),
            sanitize_field(zgloszenie.istotnosc_integr),
            sanitize_field(zgloszenie.istotnosc_dost),
            sanitize_field(zgloszenie.zglaszajacy_podpis),
            sanitize_field(zgloszenie.kierownik_lim_opinia),
            sanitize_field(zgloszenie.kierownik_lim_podpis),
            sanitize_field(zgloszenie.kierownik_lim_data),
            sanitize_field(zgloszenie.kierownik_km_opinia),
            sanitize_field(zgloszenie.kierownik_km_podpis),
            sanitize_field(zgloszenie.kierownik_km_data),
            sanitize_field(zgloszenie.realizacja_opis),
            sanitize_field(zgloszenie.realizacja_podpis),
            sanitize_field(zgloszenie.realizacja_data),
            sanitize_field(zgloszenie.potwierdzenie_podpis),
            sanitize_field(zgloszenie.potwierdzenie_data),
            urzadzenie_id,
            sanitize_field(zgloszenie.notatki)
        ])
    return response


def export_system_operacyjny_to_csv(modeladmin, request, queryset):
    response = HttpResponse(content_type='text/csv')
    now = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    filename = f'systemy_operacyjne_{now}.csv'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    writer = csv.writer(response, quotechar='"', quoting=csv.QUOTE_MINIMAL)
    writer.writerow([
        'Typ Systemu Operacyjnego'
    ])

    systemy_operacyjne = queryset.objects.all()
    for system_operacyjny in systemy_operacyjne:
        writer.writerow([
            system_operacyjny.typ_system_operacyjny.replace('\n', ' ').replace('\r', ' ')
        ])
    return response


def generate_docx(modeladmin, request, queryset):
    # przekształcenie danych w format odpowiedni do wydruku
    ids = [obj.id for obj in queryset]
    # pobranie danych z bazy danych
    with connection.cursor() as cursor:
        cursor.execute(f"SELECT * FROM bazadanych_zgloszenie WHERE id IN ({','.join(map(str, ids))})")
        rows = cursor.fetchall()
        columns = [col[0] for col in cursor.description]
        df = pd.DataFrame(rows, columns=columns)
    buffer = BytesIO()
    # zapisywanie danych z bazy danych do odpowiednich komórek pliku formatki
    for index, row in df.iterrows():
        # TA FORMATKA MUSI BYĆ WE WSKAZANEJ LOKALIZACJI Z WSKAZANĄ NAZWĄ!!!
        # KAŻDA ZMIANA FORMATKI BĘDZIE SKUTKOWAŁA ZMIANĄ PONIŻSZEGO KODU!
        # Poniższa wersja działa tylko dla tej konkretnej wersji
        document = Document("bazadanych/Zgloszenie_szablon.docx")
        table1 = document.tables[0]
        table2 = document.tables[1]
        # Wpisywanie danych do pierwszej tabeli formatki (nagłówka)
        table1.cell(0, 3).paragraphs[0].text = row['nr_zgloszenia'] or ""
        table1.cell(0, 3).add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        table1.cell(0, 3).paragraphs[1].text = row['nr_EZD'] or ""
        # Wpisywanie dat wymaga zastosowania takiego stwierdzenia "if",
        # by uniknąć wstawienia "0" zamiast pustego miejsca przy braku podanej daty
        if str(row['data_zgloszenia']) == "None":
            table1.cell(1, 3).paragraphs[-1].text = ""
        else:
            table1.cell(1, 3).paragraphs[-1].text = str(row['data_zgloszenia'])
        # Wpisywanie danych do pól pomiędzy tabelami
        document.paragraphs[2].add_run(row['nazwa_zakladu'] or "")
        document.paragraphs[3].add_run(row['laboratorium'] or "")
        document.paragraphs[4].add_run(row['zglaszajacy'] or "")
        document.paragraphs[5].add_run(row['telefon'] or "")
        document.paragraphs[6].add_run(row['nr_pomieszczenia'] or "")

        # Wpisywanie danych do drugiej tabeli w formatce
        table2.cell(0, 0).paragraphs[-1].add_run(row['nazwa_urzadzenia'] or "")
        table2.cell(1, 0).paragraphs[-1].add_run(row['dostep_do_sieci'] or "")
        table2.cell(2, 0).paragraphs[-1].add_run(row['nr_ewidencyjny'] or "")
        table2.cell(2, 2).paragraphs[-1].add_run(row['nr_gniazdka_lan'] or "")
        table2.cell(3, 0).paragraphs[-1].add_run(row['opis_zgloszenia'] or "")
        table2.cell(4, 0).paragraphs[-1].add_run(row['oczekiwania'] or "")
        table2.cell(5, 3).text = row['istotnosc_pouf'] or ""
        table2.cell(6, 3).text = row['istotnosc_integr'] or ""
        table2.cell(7, 3).text = row['istotnosc_dost'] or ""
        table2.cell(9, 0).paragraphs[0].text = row['zglaszajacy_podpis'] or ""
        table2.cell(10, 0).paragraphs[1].add_run(row['kierownik_lim_opinia'] or "")
        table2.cell(11, 0).paragraphs[0].text = row['kierownik_lim_podpis'] or ""
        if str(row['kierownik_lim_data']) == "None":
            table2.cell(11, 0).paragraphs[2].text = ""
        else:
            table2.cell(11, 0).paragraphs[2].text = str(row['kierownik_lim_data'])

        table2.cell(12, 0).paragraphs[1].add_run(row['kierownik_km_opinia'] or "")
        table2.cell(13, 0).paragraphs[0].add_run(row['kierownik_km_podpis'] or "")
        if str(row['kierownik_km_data']) == "None":
            table2.cell(13, 0).paragraphs[-1].text = ""
        else:
            table2.cell(13, 0).paragraphs[-1].text = str(row['kierownik_km_data'])
        table2.cell(14, 0).paragraphs[1].add_run(row['realizacja_opis'] or "")
        table2.cell(15, 0).paragraphs[0].text = row['realizacja_podpis'] or ""
        if str(row['realizacja_data']) == "None":
            table2.cell(15, 0).paragraphs[-1].text = ""
        else:
            table2.cell(15, 0).paragraphs[-1].text = str(row['realizacja_data'])

        table2.cell(16, 0).paragraphs[2].add_run(row['potwierdzenie_podpis'] or "")
        if str(row['potwierdzenie_data']) == "None":
            table2.cell(16, 0).paragraphs[-1].text = ""
        else:
            table2.cell(16, 0).paragraphs[-1].text = str(row['potwierdzenie_data'])

        # Przygotowanie nazwy pliku (usuwanie z niej znaków zabronionych i dodanie odpowiedniego sufixa)
        prefix = (row['nr_zgloszenia'] or "Wybierz-kolego-numer-zgloszenia-nastepnym-razem-czy-cos").replace(".", "n").replace("/", "n")
        filename = f"{prefix}-F1-IP003-A-IT Zgloszenie-pomocy-technicznej-systemow-informatyki-metrologicznej-(SIM).docx"
        cleaned_name = re.sub(r'[\\/*?:"<>|]', "", filename)
        # Zapisanie pliku
        document.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={cleaned_name}'

        return response


# New view functions
def view_urzadzenia(modeladmin, request, queryset):
    selected = queryset.values_list('urzadzenie_id', flat=True)
    url = reverse('admin:view_urzadzenia', args=[','.join(map(str, selected))])
    return redirect(url)


def view_urzadzenia_view(request, urzadzenie_ids):
    urzadzenie_ids = urzadzenie_ids.split(',')
    urzadzenia = Urzadzenie.objects.filter(pim_id__in=urzadzenie_ids)
    for urzadzenie in urzadzenia:
        urzadzenie.admin_url = reverse('admin:bazadanych_urzadzenie_change', args=[urzadzenie.pim_id])
    context = dict(
        admin.site.each_context(request),
        title="Associated Urządzenia",
        urzadzenia=urzadzenia,
    )
    return TemplateResponse(request, "admin/view_urzadzenia.html", context)


def view_zgloszenia(modeladmin, request, queryset):
    selected = queryset.values_list('pim_id', flat=True)
    url = reverse('admin:view_zgloszenia', args=[','.join(map(str, selected))])
    return redirect(url)


def view_zgloszenia_view(request, urzadzenie_ids):
    urzadzenie_ids = urzadzenie_ids.split(',')
    zgloszenia = Zgloszenie.objects.filter(urzadzenie_id__in=urzadzenie_ids)
    for zgloszenie in zgloszenia:
        zgloszenie.admin_url = reverse('admin:bazadanych_zgloszenie_change', args=[zgloszenie.id])
    context = dict(
        admin.site.each_context(request),
        title="Associated Zgloszenia",
        zgloszenia=zgloszenia,
    )
    return TemplateResponse(request, "admin/view_zgloszenia.html", context)


generate_docx.short_description = "Wygeneruj formularz w formacie .docx"
export_zgloszenie_to_csv.short_description = "Wyeksportuj zgłoszenia do csv"
export_urzadzenie_to_csv.short_description = "Wyeksportuj urządzenia do csv"
export_system_operacyjny_to_csv.short_description = "Wyeksportuj systemy operacyjne do csv"
view_urzadzenia.short_description = "Wyświetl powiązane Urządzenia"
view_zgloszenia.short_description = "Wyświetl powiązane Zgłoszenia"
