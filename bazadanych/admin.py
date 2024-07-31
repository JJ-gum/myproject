from django.contrib import admin
from django.http import HttpResponse
from django.db import connection
from .models import Zgloszenie, Urzadzenie, SystemOperacyjny
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import pandas as pd


# Funkcja generująca dokument .docx.
# WYMAGA SZABLONU W ODPOWIEDNIM MIEJSCU
def generate_docx(modeladmin, request, queryset):
    # przekształcenie danych w format odpowiedni do wydruku
    ids = [obj.id for obj in queryset]
    # pobranie danych z bazy danych
    with connection.cursor() as cursor:
        cursor.execute(f"SELECT * FROM bazadanych_zgloszenie WHERE id IN ({','.join(map(str, ids))})")
        rows = cursor.fetchall()
        columns = [col[0] for col in cursor.description]
        df = pd.DataFrame(rows, columns=columns)
    # zapisywanie danych z bazy danych do odpowiednich komórek pliku formatki
    for index, row in df.iterrows():
        # TA FORMATKA MUSI BYĆ WE WSKAZANEJ LOKALIZACJI Z WSKAZANĄ NAZWĄ!!!
        # KAŻDA ZMIANA FORMATKI BĘDZIE SKUTKOWAŁA ZMIANĄ PONIŻSZEGO KODU!
        # Poniższa wersja działa tylko dla tej konkretnej wersji
        document = Document("bazadanych/Zgloszenie_szablon.docx")
        table1 = document.tables[0]
        table2 = document.tables[1]
        # Wpisywanie danych do pierwszej tabeli formatki (nagłówka)
        table1.cell(0, 3).paragraphs[0].text = row['nr_zgloszenia'] or ""
        table1.cell(0, 3).add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        table1.cell(0, 3).paragraphs[1].text = row['nr_EZD'] or ""
        # Wpisywanie dat wymaga zastosowania takiego stwierdzenia "if",
        # by uniknąć wstawienia "0" zamiast pustego miejsca przy braku podanej daty
        if str(row['data_zgloszenia']) == "None":
            table1.cell(1, 3).paragraphs[-1].text = ""
        else:
            table1.cell(1, 3).paragraphs[-1].text = str(row['data_zgloszenia'])
        # Wpisywanie danych do pól pomiędzy tabelami
        document.paragraphs[2].add_run(row['nazwa_zakladu'] or "")
        document.paragraphs[3].add_run(row['laboratorium'] or "")
        document.paragraphs[4].add_run(row['zglaszajacy'] or "")
        document.paragraphs[5].add_run(row['telefon'] or "")
        document.paragraphs[6].add_run(row['nr_pomieszczenia'] or "")

        # Wpisywanie danych do drugiej tabeli w formatce
        table2.cell(0, 0).paragraphs[-1].add_run(row['nazwa_urzadzenia'] or "")
        table2.cell(1, 0).paragraphs[-1].add_run(row['dostep_do_sieci'] or "")
        table2.cell(2, 0).paragraphs[-1].add_run(row['nr_ewidencyjny'] or "")
        table2.cell(2, 2).paragraphs[-1].add_run(row['nr_gniazdka_lan'] or "")
        table2.cell(3, 0).paragraphs[-1].add_run(row['opis_zgloszenia'] or "")
        table2.cell(4, 0).paragraphs[-1].add_run(row['oczekiwania'] or "")
        table2.cell(5, 3).text = row['istotnosc_pouf'] or ""
        table2.cell(6, 3).text = row['istotnosc_integr'] or ""
        table2.cell(7, 3).text = row['istotnosc_dost'] or ""
        table2.cell(9, 0).paragraphs[0].text = row['zglaszajacy_podpis'] or ""
        table2.cell(10, 0).paragraphs[1].add_run(row['kierownik_lim_opinia'] or "")
        table2.cell(11, 0).paragraphs[0].text = row['kierownik_lim_podpis'] or ""
        if str(row['kierownik_lim_data']) == "None":
            table2.cell(11, 0).paragraphs[2].text = ""
        else:
            table2.cell(11, 0).paragraphs[2].text = str(row['kierownik_lim_data'])

        table2.cell(12, 0).paragraphs[1].add_run(row['kierownik_km_opinia'] or "")
        table2.cell(13, 0).paragraphs[0].add_run(row['kierownik_km_podpis'] or "")
        if str(row['kierownik_km_data']) == "None":
            table2.cell(13, 0).paragraphs[-1].text = ""
        else:
            table2.cell(13, 0).paragraphs[-1].text = str(row['kierownik_km_data'])
        table2.cell(14, 0).paragraphs[1].add_run(row['realizacja_opis'] or "")
        table2.cell(15, 0).paragraphs[-1].text = row['realizacja_podpis'] or ""
        if str(row['realizacja_data']) == "None":
            table2.cell(15, 0).paragraphs[-1].text = ""
        else:
            table2.cell(15, 0).paragraphs[-1].text = str(row['realizacja_data'])

        table2.cell(16, 0).paragraphs[2].add_run(row['potwierdzenie_podpis'] or "")
        if str(row['potwierdzenie_data']) == "None":
            table2.cell(16, 0).paragraphs[-1].text = ""
        else:
            table2.cell(16, 0).paragraphs[-1].text = str(row['potwierdzenie_data'])

        # Przygotowanie nazwy pliku (usuwanie z niej znaków zabronionych i dodanie odpowiedniego sufixa)
        prefix = (row['nr_zgloszenia'] or "Wybierz-kolego-numer-zgloszenia-nastepnym-razem-czy-cos").replace(".", "n").replace("/", "n")
        filename = f"{prefix}-F1-IP003-A-IT Zgloszenie-pomocy-technicznej-systemow-informatyki-metrologicznej-(SIM).docx"
        cleaned_name = re.sub(r'[\\/*?:"<>|]', "", filename)
        # Zapisanie pliku
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={cleaned_name}'

        return response


# Opis funkcji generującej dokument — napis wyświetlany przy wybraniu tej opcji w panelu administratora
generate_docx.short_description = "Wygeneruj formularz w formacie .docx"


# Ustawienia dotyczące wyświetlania Zgłoszenia w panelu administratora
@admin.register(Zgloszenie)
class Zgloszenie(admin.ModelAdmin):
    list_display = ('id', 'nr_EZD_ID_koszulki', 'nr_zgloszenia', 'data_zgloszenia', 'nazwa_zakladu', 'laboratorium', 'urzadzenie_id',)
    search_fields = ('nr_zgloszenia', 'data_zgloszenia', 'nazwa_zakladu', 'laboratorium', 'nr_EZD_ID_koszulki', 'urzadzenie_id')
    list_filter = ('data_zgloszenia', 'nazwa_zakladu', 'laboratorium',)
    ordering = ('-id',)
    actions = [generate_docx]


# Ustawienia dotyczące wyświetlania Urządzenia w panelu administratora
@admin.register(Urzadzenie)
class Urzadzenie(admin.ModelAdmin):
    list_display = ('pim_id', 'data_rejestracji', 'laboratorium', 'numer_ewidencyjny')
    search_fields = ('pim_id', 'data_rejestracji', 'numer_ewidencyjny')
    list_filter = ('laboratorium',)
    ordering = ('-pim_id',)


# Ustawienia dotyczące wyświetlania Systemu operacyjnego w panelu administratora
admin.site.register(SystemOperacyjny)
